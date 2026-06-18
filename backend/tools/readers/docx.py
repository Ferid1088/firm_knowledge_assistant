"""DOCXReaderTool — extract text, sections, and tables from DOCX files."""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from backend.core.tool_base import ToolMetadata
from backend.tools.readers.base import FileReaderTool, RawContent

logger = logging.getLogger(__name__)


class DOCXReaderTool(FileReaderTool):
    """Extract paragraphs, headings, and tables from DOCX files via python-docx."""

    metadata = ToolMetadata(
        name="reader:docx",
        version="1.0.0",
        tool_type="reader",
        description="Extract text, tables, and structure from DOCX files",
        dependencies=["python-docx>=0.8.11"],
        capabilities={
            "formats": [".docx", ".DOCX"],
            "max_size": 50 * 1024 * 1024,
            "supports_tables": True,
            "supports_sections": True,
        },
        is_production=True,
        is_async=True,
    )

    supported_extensions = [".docx", ".DOCX"]
    format_name = "docx"

    def _validate_dependencies(self) -> None:
        """Check that python-docx is installed (pip name differs from import name)."""
        # python-docx installs as 'python-docx' but imports as 'docx'
        import importlib
        try:
            importlib.import_module("docx")
        except ImportError:
            raise ImportError(
                f"Tool '{self.metadata.name}' requires 'python-docx>=0.8.11'. "
                "Install with: uv pip install python-docx"
            )

    async def execute(self, input_data: str, **kwargs) -> RawContent:
        """Read the DOCX file and return paragraphs, sections, and table data."""
        file_path = str(input_data)
        is_valid, error = await self.validate_input(file_path)
        if not is_valid:
            return RawContent(
                format="docx",
                content="",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                is_corrupted=True,
                warnings=[error],
            )

        try:
            from docx import Document

            doc = Document(file_path)
            parts = []
            sections = []
            tables_data = []

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                parts.append(text)
                if para.style.name.startswith("Heading"):
                    level_char = para.style.name.split()[-1]
                    level = int(level_char) if level_char.isdigit() else 1
                    sections.append({"level": level, "text": text, "position": i})

            for table in doc.tables:
                rows_data = [[cell.text for cell in row.cells] for row in table.rows]
                tables_data.append({
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": rows_data,
                })
                parts.append("--- Table ---")
                for row in rows_data:
                    parts.append(" | ".join(row))

            # Use non-empty paragraph count as page_count proxy (spec §5.3)
            page_count = len([p for p in doc.paragraphs if p.text.strip()])

            p = Path(file_path)
            return RawContent(
                format="docx",
                content="\n".join(parts),
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                page_count=page_count,
                table_count=len(tables_data),
                tables=tables_data,
                sections=sections,
                file_size=p.stat().st_size,
                modified_time=datetime.fromtimestamp(p.stat().st_mtime).isoformat(),
            )

        except Exception as e:
            return self._handle_error(e, file_path)
